"""議事録の Markdown / プレーンテキストを Word(.docx) / Excel(.xlsx) に変換する。

100%版:
- Word: ヘッダー/フッター・基本情報を表化・見出し自動採番・ページ番号・表紙ブロック
- Excel: 5シート分割（カバー / 基本情報 / 議論詳細 / 決定事項 / 文字起こし）
        + 決定事項を「No / 内容 / 担当 / 期限 / 状況」のタスク表に + 自動フィルタ + 印刷設定
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional


# ===========================================================================
# パーサ: Markdown 議事録 → 構造化データ
# ===========================================================================

_BULLET_RE = re.compile(r"^\s*(?:[-*・●○◯▪■]|\d+[.)])\s+(.*)$")
_HEADING_RE = re.compile(r"^\s*(#{1,4})\s+(.*)$")
_NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+)\.\s*([^\s].*)$")
_KEYVALUE_RE = re.compile(r"^([^：:]{1,30})\s*[：:]\s*(.+)$")
_AGENDA_RE = re.compile(r"^[\[【]?(議題|案件|テーマ)\s*\d*[\]】]?\s*[:：]?\s*(.*)$")
_MD_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$")


def _parse_md_table_row(line: str) -> list[str]:
    """| a | b | c | のような行をパースしてセル配列を返す（前後のパイプは除去）"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    # <br> をセル内改行として保持しつつ分割
    cells = [c.strip() for c in s.split("|")]
    # 各セル内の <br> を実際の改行に変換
    cells = [c.replace("<br>", "\n").replace("<BR>", "\n") for c in cells]
    return cells


@dataclass
class Section:
    """議事録の意味のあるセクション。"""
    title: str
    level: int = 1               # 1=大見出し, 2=中, 3=小
    info_pairs: list[tuple[str, str]] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)  # 各表は2D配列
    subsections: list["Section"] = field(default_factory=list)

    def is_basic_info(self) -> bool:
        keys = {"会議名", "日時", "場所", "出席", "出席者", "欠席",
                "欠席者", "議事録作成", "議事録担当", "司会"}
        return any(k in self.title for k in ("基本情報", "概要")) or \
               (bool(self.info_pairs) and not self.bullets and not self.paragraphs and
                any(k.replace("：", "").strip() in keys for k, _ in self.info_pairs))

    def is_decisions(self) -> bool:
        return any(k in self.title for k in ("決定事項", "Action", "アクション", "ToDo", "決定", "決まった"))

    def is_agenda(self) -> bool:
        return any(k in self.title for k in ("議論詳細", "議論", "議題", "本件", "内容", "議事"))


@dataclass
class MeetingMeta:
    """UIから渡される会議メタ情報。"""
    company: str = "ニッケン建設株式会社"
    template_name: str = "議事録"
    meeting_date: str = ""
    meeting_location: str = ""
    attendees_note: str = ""
    generated_at: datetime = field(default_factory=datetime.now)


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def parse_minutes(markdown_text: str) -> tuple[Optional[str], list[Section]]:
    """
    議事録 Markdown を「タイトル + セクション一覧」に変換。
    ・# 見出し → Section(level=1)
    ・## or 「1. 基本情報」など → Section(level=1 か 2 を文脈で決定)
    ・本文中の「会議名： xxx」は info_pairs へ
    ・「・xxx」は bullets へ
    """
    lines = markdown_text.splitlines()
    title: Optional[str] = None
    sections: list[Section] = []
    current: Optional[Section] = None
    section_stack: list[Section] = []

    def push_section(sec: Section) -> None:
        nonlocal current
        if sec.level <= 1:
            sections.append(sec)
            section_stack.clear()
            section_stack.append(sec)
        else:
            # ネスト
            while section_stack and section_stack[-1].level >= sec.level:
                section_stack.pop()
            if section_stack:
                section_stack[-1].subsections.append(sec)
            else:
                sections.append(sec)
            section_stack.append(sec)
        current = sec

    # Markdown パイプ表を検出するため、複数行をまとめる必要がある
    # 一度全体を走査して、表の連続を検出してから他の処理に進む
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # ===== Markdown パイプ表の検出 =====
        # 連続する table row（途中にseparatorを含む）を1つの表として収集
        if _MD_TABLE_ROW_RE.match(line):
            table_rows: list[list[str]] = []
            j = i
            while j < len(lines):
                lj = lines[j].rstrip()
                if not lj.strip():
                    break
                if _MD_TABLE_SEP_RE.match(lj):
                    j += 1
                    continue
                if _MD_TABLE_ROW_RE.match(lj):
                    table_rows.append(_parse_md_table_row(lj))
                    j += 1
                else:
                    break
            if len(table_rows) >= 2:  # ヘッダー＋1行以上で表とみなす
                if current is None:
                    current = Section(title="")
                    sections.append(current)
                    section_stack.append(current)
                current.tables.append(table_rows)
                i = j
                continue
            # 表として認識できない単独行は通常の段落として扱う

        # # 見出し
        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = _strip_inline_md(m.group(2))
            if level == 1 and title is None and current is None:
                title = text
                i += 1
                continue
            push_section(Section(title=text, level=min(level, 3)))
            i += 1
            continue

        # 「1. 基本情報」等の番号付きトップ見出し
        m = _NUMBERED_HEADING_RE.match(line)
        if m and len(m.group(2)) <= 30 and "：" not in m.group(2) and ":" not in m.group(2):
            text = _strip_inline_md(f"{m.group(1)}. {m.group(2)}")
            push_section(Section(title=text, level=1))
            i += 1
            continue

        # 箇条書き
        m = _BULLET_RE.match(line)
        if m:
            if current is None:
                current = Section(title="")
                sections.append(current)
                section_stack.append(current)
            current.bullets.append(_strip_inline_md(m.group(1)))
            i += 1
            continue

        # キー：値
        m = _KEYVALUE_RE.match(stripped)
        if m and len(m.group(1).strip()) <= 20:
            if current is None:
                current = Section(title="")
                sections.append(current)
                section_stack.append(current)
            current.info_pairs.append((
                _strip_inline_md(m.group(1)),
                _strip_inline_md(m.group(2)),
            ))
            i += 1
            continue

        # それ以外は段落
        if current is None:
            current = Section(title="")
            sections.append(current)
            section_stack.append(current)
        current.paragraphs.append(_strip_inline_md(stripped))
        i += 1

    return title, sections


# ===========================================================================
# 決定事項のタスク化（Excel用）
# ===========================================================================

@dataclass
class TaskRow:
    no: int
    content: str
    owner: str = ""
    deadline: str = ""
    status: str = "未着手"


# 「担当：〇〇」明示型
_OWNER_EXPLICIT_RE = re.compile(r"(?:担当|責任者|担当者)[：:\s]*([^\s、,。]+)")

# 役職パターン（社内テンプレート想定）
_ROLE_PATTERNS = [
    r"代表", r"社長", r"取締役", r"執行役員",
    r"部長", r"次長", r"課長", r"係長", r"主任", r"所長", r"工場長",
    r"先生", r"さん", r"様",
]
# 「〇〇代表が／〇〇部長は／中辻良太が／高橋部長より」を拾う
# 左境界（行頭・記号・句読点・括弧 + 助詞「を/に/で/と/から/より」）
# → 人名候補（漢字/カタカナのみ・1〜6字）
# → 役職パターン
# → 終端助詞（が/は/より/から/に）
_LEFT_BOUNDARY = r"(?:^|[\s、。\.,「『（(【\[／/をにでとは・]|から)"
_OWNER_NAME_RE = re.compile(
    _LEFT_BOUNDARY +
    r"([一-龯々ァ-ヴー]{1,6}(?:" + "|".join(_ROLE_PATTERNS) + r"))"
    r"(?:が|は|より|から|に[はが]?|を中心)"
)
# フルネーム単独（漢字2〜4 + 空白？ + 漢字1〜4 ＋ 助詞）
_FULLNAME_RE = re.compile(
    _LEFT_BOUNDARY +
    r"([一-龯々]{2,4}\s?[一-龯々]{1,4})\s*(?:が|は|より)"
)

# 期限パターン（順番が重要：具体的 → 曖昧）
_DEADLINE_PATTERNS = [
    re.compile(r"(\d{4}\s*[/年\-]\s*\d{1,2}\s*[/月\-]\s*\d{1,2}\s*日?)"),
    re.compile(r"(\d{1,2}\s*[/月]\s*\d{1,2}\s*日?)(?:まで|までに|期限|期日)"),
    re.compile(r"(\d{1,2}\s*月\s*\d{1,2}\s*日)"),
    re.compile(r"(来週\s*[月火水木金土日](?:曜日?)?)"),
    re.compile(r"(今週\s*[月火水木金土日](?:曜日?)?)"),
    re.compile(r"(今週末|来週末|今月末|来月末|月末|期末|今月中|来月中)"),
    re.compile(r"(来週|今週|来月|今月|翌月|次回(?:会議)?)"),
]


def _looks_like_real_owner(candidate: str) -> bool:
    """役職っぽい候補を、「までに」「ことから」等の偽物から弾く。"""
    if not candidate:
        return False
    bad_substrings = ("まで", "までに", "ことか", "として", "ため",
                      "うえ", "あと", "もとに", "について")
    return not any(b in candidate for b in bad_substrings)


def extract_tasks(decisions: list[str]) -> list[TaskRow]:
    """決定事項の文字列リストから、担当・期限を推測してタスク表に変換。"""
    rows = []
    for i, raw in enumerate(decisions, start=1):
        text = raw.strip()
        # 行頭の箇条書き記号や番号を取り除く
        text = re.sub(r"^[・●○◯▪■\-*]\s*", "", text)
        text = re.sub(r"^\d+[.)]\s*", "", text)

        owner = ""
        deadline = ""

        # 期限: より具体的なパターンから順に（先に抽出して取り除く）
        for pat in _DEADLINE_PATTERNS:
            m = pat.search(text)
            if m:
                deadline = m.group(1).strip()
                break

        # 担当探索用に、抽出した期限とその前後の助詞を取り除く
        # 例: 「来週月曜までに」「5月15日までに」を text から消す
        # 重要：選択肢は長いものから順に並べる（まで | までに ではなく までに | まで）
        text_for_owner = text
        if deadline:
            text_for_owner = re.sub(
                re.escape(deadline) + r"(?:までには|までに|まで|期限内|期限まで|期日まで|期日)?",
                " ",
                text_for_owner,
            )

        # 担当: 明示優先 → 役職付き名前 → フルネーム
        m = _OWNER_EXPLICIT_RE.search(text_for_owner)
        if m and _looks_like_real_owner(m.group(1)):
            owner = m.group(1).strip()
        else:
            m = _OWNER_NAME_RE.search(text_for_owner)
            if m and _looks_like_real_owner(m.group(1)):
                owner = m.group(1).strip()
            else:
                m = _FULLNAME_RE.search(text_for_owner)
                if m and _looks_like_real_owner(m.group(1)):
                    owner = m.group(1).strip()

        rows.append(TaskRow(no=i, content=text, owner=owner, deadline=deadline))
    return rows


# ===========================================================================
# 共通ユーティリティ
# ===========================================================================

def _set_eastasia_font(run, font_name: str = "游ゴシック") -> None:
    """python-docx で東アジアフォントを設定する。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def _add_page_number_field(paragraph) -> None:
    """フッター段落に「ページ X / Y」のフィールドを追加する。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.add_run("ページ ")
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instr)
    run._element.append(fldChar2)

    paragraph.add_run(" / ")
    run2 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run2._element.append(fldChar3)
    run2._element.append(instr2)
    run2._element.append(fldChar4)


# ===========================================================================
# Word(.docx) エクスポート（100%版）
# ===========================================================================

def to_docx_bytes(
    markdown_text: str,
    title: str = "議事録",
    meta: Optional[MeetingMeta] = None,
) -> bytes:
    """議事録 Markdown を、印刷品質の Word ファイルに変換する。"""
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    if meta is None:
        meta = MeetingMeta()

    parsed_title, sections = parse_minutes(markdown_text)
    document_title = parsed_title or title

    doc = Document()

    # ===== ページ設定 =====
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    # ===== 既定フォント =====
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游ゴシック")
    rfonts.set(qn("w:ascii"), "游ゴシック")
    rfonts.set(qn("w:hAnsi"), "游ゴシック")

    # ===== ヘッダー（会社名 / 会議種別） =====
    header = doc.sections[0].header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_para.add_run(f"{meta.company}　{meta.template_name}")
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_eastasia_font(h_run)

    # ヘッダー下に罫線
    pPr = h_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ===== フッター（ページ番号 / 生成日） =====
    footer = doc.sections[0].footer
    # ページ番号（中央）
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(f_para)
    for run in f_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _set_eastasia_font(run)

    # 生成日（右）
    f_para2 = footer.add_paragraph(
        f"生成日: {meta.generated_at.strftime('%Y-%m-%d %H:%M')}"
    )
    f_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in f_para2.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        _set_eastasia_font(run)

    # ===== タイトル =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(document_title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    _set_eastasia_font(title_run)

    # サブタイトル（日時・場所・出席者）
    if meta.meeting_date or meta.meeting_location:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        if meta.meeting_date:
            parts.append(meta.meeting_date)
        if meta.meeting_location:
            parts.append(meta.meeting_location)
        sub_run = sub_para.add_run("　|　".join(parts))
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_eastasia_font(sub_run)

    # 区切り罫線
    sep = doc.add_paragraph()
    sep_run = sep.add_run("─" * 35)
    sep_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ===== セクション本体 =====
    section_no = 0  # 自動採番
    for section in sections:
        if not section.title and not section.info_pairs and \
           not section.bullets and not section.paragraphs:
            continue

        # 見出し（タイトルが空でなければ採番）
        if section.title:
            section_no += 1
            heading_text = section.title
            # すでに番号付きでなければ、自動付与
            if not re.match(r"^\d+[.)]", heading_text):
                heading_text = f"{section_no}. {heading_text}"
            heading = doc.add_heading(heading_text, level=1)
            for run in heading.runs:
                _set_eastasia_font(run)
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6B)

        # 基本情報セクションは表で出力（最重要の改善）
        if section.info_pairs and section.is_basic_info():
            tbl = doc.add_table(rows=len(section.info_pairs), cols=2)
            tbl.style = "Light Grid Accent 1"
            for row, (key, value) in zip(tbl.rows, section.info_pairs):
                row.cells[0].width = Cm(4.0)
                row.cells[1].width = Cm(13.0)
                # キー
                row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p1 = row.cells[0].paragraphs[0]
                p1.text = ""
                k_run = p1.add_run(key)
                k_run.bold = True
                _set_eastasia_font(k_run)
                # シェーディング
                tcPr = row.cells[0]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EAF1FA")
                tcPr.append(shd)
                # 値
                row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p2 = row.cells[1].paragraphs[0]
                p2.text = ""
                v_run = p2.add_run(value)
                _set_eastasia_font(v_run)
            doc.add_paragraph("")
        else:
            # 基本情報以外でも kv はあれば段落で出す
            for key, value in section.info_pairs:
                p = doc.add_paragraph()
                k_run = p.add_run(f"{key}：")
                k_run.bold = True
                _set_eastasia_font(k_run)
                v_run = p.add_run(value)
                _set_eastasia_font(v_run)

        # 段落
        for para_text in section.paragraphs:
            p = doc.add_paragraph(para_text)
            for run in p.runs:
                _set_eastasia_font(run)

        # 箇条書き
        for bullet in section.bullets:
            try:
                p = doc.add_paragraph(bullet, style="List Bullet")
            except KeyError:
                p = doc.add_paragraph(f"・{bullet}")
            for run in p.runs:
                _set_eastasia_font(run)

        # ===== Markdownパイプ表 → Word表 =====
        for table_rows in section.tables:
            if not table_rows:
                continue
            n_cols = max(len(r) for r in table_rows)
            n_rows = len(table_rows)
            tbl = doc.add_table(rows=n_rows, cols=n_cols)
            tbl.style = "Light Grid Accent 1"
            for r_idx, row_cells in enumerate(table_rows):
                for c_idx in range(n_cols):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    val = row_cells[c_idx] if c_idx < len(row_cells) else ""
                    # セル内に複数行（<br>で分けた行）がある場合に対応
                    cell.text = ""
                    p_first = cell.paragraphs[0]
                    lines_in_cell = val.split("\n")
                    for li, ln in enumerate(lines_in_cell):
                        para = p_first if li == 0 else cell.add_paragraph()
                        run = para.add_run(ln)
                        _set_eastasia_font(run)
                        if r_idx == 0:
                            run.bold = True
                            # ヘッダー行は薄い青で塗る
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "DDEBF7")
                            tcPr.append(shd)
                    # 章タイトル行（同じ内容が3列に並ぶ）の検出
                    if r_idx > 0 and c_idx == 0 and n_cols >= 3:
                        if (
                            len(row_cells) >= 3
                            and row_cells[0] == row_cells[1] == row_cells[2]
                            and row_cells[0]
                        ):
                            # 章タイトル行は薄いグレーで塗る
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "EAEAEA")
                            tcPr.append(shd)
                            for run in p_first.runs:
                                run.bold = True
            doc.add_paragraph("")

        # サブセクション
        for sub in section.subsections:
            sub_heading = doc.add_heading(sub.title, level=2)
            for run in sub_heading.runs:
                _set_eastasia_font(run)
                run.font.color.rgb = RGBColor(0x33, 0x55, 0x88)
            for key, value in sub.info_pairs:
                p = doc.add_paragraph()
                k_run = p.add_run(f"{key}：")
                k_run.bold = True
                _set_eastasia_font(k_run)
                v_run = p.add_run(value)
                _set_eastasia_font(v_run)
            for para_text in sub.paragraphs:
                p = doc.add_paragraph(para_text)
                for run in p.runs:
                    _set_eastasia_font(run)
            for bullet in sub.bullets:
                try:
                    p = doc.add_paragraph(bullet, style="List Bullet")
                except KeyError:
                    p = doc.add_paragraph(f"・{bullet}")
                for run in p.runs:
                    _set_eastasia_font(run)
            # サブセクション内の表も処理
            for table_rows in sub.tables:
                if not table_rows:
                    continue
                n_cols = max(len(r) for r in table_rows)
                n_rows = len(table_rows)
                tbl = doc.add_table(rows=n_rows, cols=n_cols)
                tbl.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(table_rows):
                    for c_idx in range(n_cols):
                        cell = tbl.rows[r_idx].cells[c_idx]
                        val = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell.text = ""
                        p_first = cell.paragraphs[0]
                        for li, ln in enumerate(val.split("\n")):
                            para = p_first if li == 0 else cell.add_paragraph()
                            run = para.add_run(ln)
                            _set_eastasia_font(run)
                            if r_idx == 0:
                                run.bold = True
                doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# Excel(.xlsx) エクスポート（100%版・マルチシート）
# ===========================================================================

def to_xlsx_bytes(
    markdown_text: str,
    title: str = "議事録",
    meta: Optional[MeetingMeta] = None,
    transcript: Optional[str] = None,
) -> bytes:
    """
    議事録 Markdown を、業務で使える Excel ファイルに変換する。

    シート構成:
      1. カバー         … タイトル・サマリ
      2. 基本情報       … 会議名/日時/場所/出席者を整理
      3. 議論詳細       … 議題ごとに行
      4. 決定事項       … No / 内容 / 担当 / 期限 / 状況 のタスク表
      5. 文字起こし     … （transcript が渡された場合のみ）
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.table import Table, TableStyleInfo

    if meta is None:
        meta = MeetingMeta()

    parsed_title, sections = parse_minutes(markdown_text)
    document_title = parsed_title or title

    # ===== 共通スタイル =====
    JP = "游ゴシック"
    title_font = Font(name=JP, size=20, bold=True, color="1F3A6B")
    subtitle_font = Font(name=JP, size=11, color="555555")
    sec_font = Font(name=JP, size=14, bold=True, color="FFFFFF")
    sec_fill = PatternFill(start_color="1F3A6B", end_color="1F3A6B", fill_type="solid")
    head_font = Font(name=JP, size=11, bold=True, color="FFFFFF")
    head_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
    body_font = Font(name=JP, size=10)
    bold_body_font = Font(name=JP, size=10, bold=True)
    info_key_fill = PatternFill(
        start_color="EAF1FA", end_color="EAF1FA", fill_type="solid"
    )
    wrap_top_left = Alignment(wrap_text=True, vertical="top", horizontal="left")
    wrap_center = Alignment(wrap_text=True, vertical="center", horizontal="center")
    wrap_top_left_indent = Alignment(
        wrap_text=True, vertical="top", horizontal="left", indent=1
    )
    thin = Side(border_style="thin", color="B4B4B4")
    cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    wb = Workbook()
    # 既定の Sheet を削除
    wb.remove(wb.active)

    def _setup_print(ws, fit_to_width: bool = True) -> None:
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_setup.fitToWidth = 1 if fit_to_width else 0
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_options.horizontalCentered = True
        ws.page_margins.left = 0.5
        ws.page_margins.right = 0.5
        ws.page_margins.top = 0.7
        ws.page_margins.bottom = 0.7
        ws.oddHeader.right.text = f"{meta.company}　{meta.template_name}"
        ws.oddHeader.right.size = 9
        ws.oddHeader.right.color = "808080"
        ws.oddFooter.center.text = "Page &P / &N"
        ws.oddFooter.right.text = (
            f"生成: {meta.generated_at.strftime('%Y-%m-%d %H:%M')}"
        )
        ws.oddFooter.right.size = 8
        ws.oddFooter.right.color = "AAAAAA"

    # =====================================================================
    # シート1: カバー
    # =====================================================================
    ws1 = wb.create_sheet("カバー")
    _setup_print(ws1, fit_to_width=True)
    ws1.row_dimensions[2].height = 40
    ws1.merge_cells("B2:F2")
    ws1["B2"] = document_title
    ws1["B2"].font = title_font
    ws1["B2"].alignment = wrap_center

    ws1.row_dimensions[3].height = 25
    ws1.merge_cells("B3:F3")
    sub = []
    if meta.meeting_date:
        sub.append(meta.meeting_date)
    if meta.meeting_location:
        sub.append(meta.meeting_location)
    ws1["B3"] = "　|　".join(sub) if sub else ""
    ws1["B3"].font = subtitle_font
    ws1["B3"].alignment = wrap_center

    # 基本情報のミニ表
    ws1.merge_cells("B5:F5")
    ws1["B5"] = "■ 基本情報"
    ws1["B5"].font = Font(name=JP, size=12, bold=True, color="1F3A6B")
    row = 6
    info_pairs_for_cover: list[tuple[str, str]] = []
    if meta.meeting_date:
        info_pairs_for_cover.append(("日時", meta.meeting_date))
    if meta.meeting_location:
        info_pairs_for_cover.append(("場所", meta.meeting_location))
    if meta.attendees_note:
        info_pairs_for_cover.append(("出席者（手入力）", meta.attendees_note))
    # 議事録本文中の基本情報も拾う
    for sec in sections:
        if sec.is_basic_info():
            for k, v in sec.info_pairs:
                info_pairs_for_cover.append((k, v))
    seen_keys = set()
    for k, v in info_pairs_for_cover:
        if k in seen_keys:
            continue
        seen_keys.add(k)
        ws1.merge_cells(start_row=row, start_column=2, end_row=row, end_column=2)
        ws1.cell(row=row, column=2, value=k)
        ws1.cell(row=row, column=2).font = bold_body_font
        ws1.cell(row=row, column=2).fill = info_key_fill
        ws1.cell(row=row, column=2).alignment = wrap_top_left_indent
        ws1.cell(row=row, column=2).border = cell_border
        ws1.merge_cells(start_row=row, start_column=3, end_row=row, end_column=6)
        ws1.cell(row=row, column=3, value=v)
        ws1.cell(row=row, column=3).font = body_font
        ws1.cell(row=row, column=3).alignment = wrap_top_left_indent
        ws1.cell(row=row, column=3).border = cell_border
        ws1.row_dimensions[row].height = 22
        row += 1

    # 列幅
    ws1.column_dimensions["A"].width = 2
    ws1.column_dimensions["B"].width = 18
    for col in "CDEF":
        ws1.column_dimensions[col].width = 18

    # =====================================================================
    # シート2: 基本情報（フル）
    # =====================================================================
    ws2 = wb.create_sheet("基本情報")
    _setup_print(ws2, fit_to_width=True)
    ws2["A1"] = "基本情報"
    ws2["A1"].font = Font(name=JP, size=14, bold=True, color="1F3A6B")
    ws2.row_dimensions[1].height = 25

    headers = ["項目", "内容"]
    for col_idx, h in enumerate(headers, start=1):
        c = ws2.cell(row=2, column=col_idx, value=h)
        c.font = head_font
        c.fill = head_fill
        c.alignment = wrap_center
        c.border = cell_border
    ws2.row_dimensions[2].height = 22

    info_collected: list[tuple[str, str]] = []
    if meta.meeting_date:
        info_collected.append(("日時", meta.meeting_date))
    if meta.meeting_location:
        info_collected.append(("場所", meta.meeting_location))
    if meta.attendees_note:
        info_collected.append(("出席者（UI入力）", meta.attendees_note))
    for sec in sections:
        if sec.is_basic_info():
            for k, v in sec.info_pairs:
                info_collected.append((k, v))

    seen_keys2 = set()
    r = 3
    for k, v in info_collected:
        if k in seen_keys2:
            continue
        seen_keys2.add(k)
        ws2.cell(row=r, column=1, value=k).font = bold_body_font
        ws2.cell(row=r, column=1).fill = info_key_fill
        ws2.cell(row=r, column=1).alignment = wrap_top_left_indent
        ws2.cell(row=r, column=1).border = cell_border
        ws2.cell(row=r, column=2, value=v).font = body_font
        ws2.cell(row=r, column=2).alignment = wrap_top_left_indent
        ws2.cell(row=r, column=2).border = cell_border
        ws2.row_dimensions[r].height = 22
        r += 1

    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 90
    ws2.freeze_panes = "A3"

    # =====================================================================
    # シート3: 議論詳細
    # =====================================================================
    ws3 = wb.create_sheet("議論詳細")
    _setup_print(ws3, fit_to_width=True)
    ws3["A1"] = "議論詳細"
    ws3["A1"].font = Font(name=JP, size=14, bold=True, color="1F3A6B")
    ws3.row_dimensions[1].height = 25

    headers3 = ["セクション", "種別", "内容"]
    for col_idx, h in enumerate(headers3, start=1):
        c = ws3.cell(row=2, column=col_idx, value=h)
        c.font = head_font
        c.fill = head_fill
        c.alignment = wrap_center
        c.border = cell_border
    ws3.row_dimensions[2].height = 22

    r = 3
    for sec in sections:
        if sec.is_basic_info() or sec.is_decisions():
            continue
        if not sec.title and not sec.bullets and not sec.paragraphs and not sec.subsections:
            continue
        # セクション帯
        ws3.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
        ws3.cell(row=r, column=1, value=sec.title or "（無題セクション）").font = sec_font
        ws3.cell(row=r, column=1).fill = sec_fill
        ws3.cell(row=r, column=1).alignment = Alignment(
            horizontal="left", vertical="center", indent=1
        )
        ws3.cell(row=r, column=1).border = cell_border
        ws3.row_dimensions[r].height = 24
        r += 1
        # 段落
        for p in sec.paragraphs:
            ws3.cell(row=r, column=1, value="").border = cell_border
            ws3.cell(row=r, column=2, value="本文").font = body_font
            ws3.cell(row=r, column=2).alignment = wrap_center
            ws3.cell(row=r, column=2).border = cell_border
            ws3.cell(row=r, column=3, value=p).font = body_font
            ws3.cell(row=r, column=3).alignment = wrap_top_left
            ws3.cell(row=r, column=3).border = cell_border
            r += 1
        # 箇条書き
        for b in sec.bullets:
            ws3.cell(row=r, column=1, value="").border = cell_border
            ws3.cell(row=r, column=2, value="•").font = body_font
            ws3.cell(row=r, column=2).alignment = wrap_center
            ws3.cell(row=r, column=2).border = cell_border
            ws3.cell(row=r, column=3, value=b).font = body_font
            ws3.cell(row=r, column=3).alignment = wrap_top_left
            ws3.cell(row=r, column=3).border = cell_border
            r += 1
        # サブセクション
        for sub in sec.subsections:
            ws3.cell(row=r, column=1, value="").border = cell_border
            ws3.cell(row=r, column=2, value="小見出し").font = bold_body_font
            ws3.cell(row=r, column=2).alignment = wrap_center
            ws3.cell(row=r, column=2).border = cell_border
            ws3.cell(row=r, column=3, value=sub.title).font = bold_body_font
            ws3.cell(row=r, column=3).alignment = wrap_top_left
            ws3.cell(row=r, column=3).border = cell_border
            r += 1
            for p in sub.paragraphs:
                ws3.cell(row=r, column=1).border = cell_border
                ws3.cell(row=r, column=2, value="本文").font = body_font
                ws3.cell(row=r, column=2).alignment = wrap_center
                ws3.cell(row=r, column=2).border = cell_border
                ws3.cell(row=r, column=3, value=p).font = body_font
                ws3.cell(row=r, column=3).alignment = wrap_top_left
                ws3.cell(row=r, column=3).border = cell_border
                r += 1
            for b in sub.bullets:
                ws3.cell(row=r, column=1).border = cell_border
                ws3.cell(row=r, column=2, value="•").font = body_font
                ws3.cell(row=r, column=2).alignment = wrap_center
                ws3.cell(row=r, column=2).border = cell_border
                ws3.cell(row=r, column=3, value=b).font = body_font
                ws3.cell(row=r, column=3).alignment = wrap_top_left
                ws3.cell(row=r, column=3).border = cell_border
                r += 1

    ws3.column_dimensions["A"].width = 4
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 95
    ws3.freeze_panes = "A3"

    # =====================================================================
    # シート4: 決定事項（タスク表）
    # =====================================================================
    ws4 = wb.create_sheet("決定事項")
    _setup_print(ws4, fit_to_width=True)
    ws4["A1"] = "決定事項・アクションアイテム"
    ws4["A1"].font = Font(name=JP, size=14, bold=True, color="1F3A6B")
    ws4.row_dimensions[1].height = 25

    headers4 = ["No", "決定内容", "担当", "期限", "状況"]
    for col_idx, h in enumerate(headers4, start=1):
        c = ws4.cell(row=2, column=col_idx, value=h)
        c.font = head_font
        c.fill = head_fill
        c.alignment = wrap_center
        c.border = cell_border
    ws4.row_dimensions[2].height = 22

    decisions_collected: list[str] = []
    for sec in sections:
        if sec.is_decisions():
            decisions_collected.extend(sec.bullets)
            decisions_collected.extend(sec.paragraphs)

    tasks = extract_tasks(decisions_collected)
    if tasks:
        for i, task in enumerate(tasks, start=3):
            ws4.cell(row=i, column=1, value=task.no).font = body_font
            ws4.cell(row=i, column=1).alignment = wrap_center
            ws4.cell(row=i, column=1).border = cell_border
            ws4.cell(row=i, column=2, value=task.content).font = body_font
            ws4.cell(row=i, column=2).alignment = wrap_top_left
            ws4.cell(row=i, column=2).border = cell_border
            ws4.cell(row=i, column=3, value=task.owner).font = body_font
            ws4.cell(row=i, column=3).alignment = wrap_center
            ws4.cell(row=i, column=3).border = cell_border
            ws4.cell(row=i, column=4, value=task.deadline).font = body_font
            ws4.cell(row=i, column=4).alignment = wrap_center
            ws4.cell(row=i, column=4).border = cell_border
            ws4.cell(row=i, column=5, value=task.status).font = body_font
            ws4.cell(row=i, column=5).alignment = wrap_center
            ws4.cell(row=i, column=5).border = cell_border
        last_row = 2 + len(tasks)
        # 自動フィルタ
        ws4.auto_filter.ref = f"A2:E{last_row}"
        # 入力規則（状況）
        try:
            from openpyxl.worksheet.datavalidation import DataValidation
            dv = DataValidation(
                type="list",
                formula1='"未着手,進行中,完了,保留,中止"',
                allow_blank=True,
            )
            ws4.add_data_validation(dv)
            dv.add(f"E3:E{last_row}")
        except Exception:
            pass
    else:
        ws4.cell(row=3, column=2, value="（決定事項は記録されていません）").font = body_font
        ws4.cell(row=3, column=2).alignment = wrap_top_left

    ws4.column_dimensions["A"].width = 5
    ws4.column_dimensions["B"].width = 70
    ws4.column_dimensions["C"].width = 15
    ws4.column_dimensions["D"].width = 18
    ws4.column_dimensions["E"].width = 12
    ws4.freeze_panes = "A3"

    # =====================================================================
    # シート5: 文字起こし（任意）
    # =====================================================================
    if transcript:
        ws5 = wb.create_sheet("文字起こし")
        _setup_print(ws5, fit_to_width=True)
        ws5["A1"] = "文字起こし全文"
        ws5["A1"].font = Font(name=JP, size=14, bold=True, color="1F3A6B")
        ws5.row_dimensions[1].height = 25
        ws5.cell(row=2, column=1, value="行").font = head_font
        ws5.cell(row=2, column=1).fill = head_fill
        ws5.cell(row=2, column=1).alignment = wrap_center
        ws5.cell(row=2, column=1).border = cell_border
        ws5.cell(row=2, column=2, value="本文").font = head_font
        ws5.cell(row=2, column=2).fill = head_fill
        ws5.cell(row=2, column=2).alignment = wrap_center
        ws5.cell(row=2, column=2).border = cell_border
        ws5.row_dimensions[2].height = 22
        for i, line in enumerate(transcript.splitlines(), start=1):
            if not line.strip():
                continue
            r = i + 2
            ws5.cell(row=r, column=1, value=i).font = body_font
            ws5.cell(row=r, column=1).alignment = wrap_center
            ws5.cell(row=r, column=1).border = cell_border
            ws5.cell(row=r, column=2, value=line).font = body_font
            ws5.cell(row=r, column=2).alignment = wrap_top_left
            ws5.cell(row=r, column=2).border = cell_border
        ws5.column_dimensions["A"].width = 6
        ws5.column_dimensions["B"].width = 110
        ws5.freeze_panes = "A3"

    # 先頭シートを「カバー」に
    wb.active = 0

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
